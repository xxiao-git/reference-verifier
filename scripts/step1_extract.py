#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Step 1: Extract references from a manuscript (.docx), a plain-text list (.txt),
or a pre-extracted JSON — format-agnostic parser.

Supported reference formats: Vancouver, APA, AMA, MLA, Chicago, Harvard, and
unstructured text. The parser extracts DOI, year, title, first author, and
journal using multi-strategy heuristics — no single format assumption.

Supported inputs:
  - .docx: Locates the References section (Heading or text match) and extracts numbered refs
  - .txt:  Each line is a reference (numbered or unnumbered)
  - .json: Pre-extracted list of {"num": int, "text": str} — passed through directly

Usage:
  python step1_extract.py --input manuscript.docx --output-dir ./output/
  python step1_extract.py --input refs.txt --output-dir ./output/
  python step1_extract.py --input manuscript.docx --output refs.json
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document


# ============================================================
# Format-agnostic field extractors
# ============================================================

def extract_doi(text):
    """Extract DOI from any format — URL form or bare DOI."""
    # doi.org URL
    url_match = re.search(r'https?://(?:dx\.)?doi\.org/(10\.\d{4,}/[^\s"\'<>]+)', text, re.IGNORECASE)
    if url_match:
        return url_match.group(1).rstrip('.')
    # Bare DOI: 10.NNNN/...
    doi_match = re.search(r'\b(10\.\d{4,9}/[^\s"\'<>;,)\']+)', text)
    if doi_match:
        return doi_match.group(1).rstrip('.')
    return None


def extract_year(text):
    """Extract publication year (first 4-digit year 1900–2030)."""
    matches = re.findall(r'\b(19\d{2}|20[0-3]\d)\b', text)
    if matches:
        return int(matches[0])
    return None


def extract_title(text, year=None):
    """Extract article title — multi-strategy, format-agnostic.

    Strategies (tried in order):
      1. Quoted title (MLA, Chicago)
      2. After "et al." (Vancouver/AMA with et al.)
      3. APA/Harvard: after "(YYYY). " or "(YYYY) 'Title',"
      4. Vancouver/AMA: second-to-last ". " segment before year
      5. Fallback: text between first ". " and year/end
    """
    # --- Strategy 1: Quoted title ---
    quoted = re.search(r'[\u201c"]([^\u201d"]{15,})[\u201d"]', text)
    if quoted:
        return quoted.group(1).strip().rstrip(".")

    # --- Strategy 2: After "et al." ---
    et_al_match = re.search(r'et\s+al\.?\s+(.+)', text, re.IGNORECASE)
    if et_al_match:
        remaining = et_al_match.group(1).strip()
        if year:
            year_str = str(year)
            year_pos = remaining.find(year_str)
            if year_pos > 0:
                title_part = remaining[:year_pos].strip().rstrip(".")
                segments = title_part.rsplit(". ", 1)
                if len(segments) == 2 and len(segments[0]) > 10:
                    return segments[0].strip()
                return title_part if len(title_part) > 10 else None
        period_pos = remaining.find(". ")
        if period_pos > 10:
            return remaining[:period_pos].strip()
        return remaining.rstrip(".") if len(remaining) > 10 else None

    # --- Strategy 3: APA / Harvard ---
    if year:
        # APA: (2024). Title. Journal...
        apa_match = re.search(r'\(\d{4}\)\.?\s+(.+?)(?:\.\s+|$)', text)
        if apa_match:
            candidate = apa_match.group(1).strip().rstrip(".")
            if len(candidate) > 10:
                return candidate
        # Harvard: (2024) 'Title',
        harvard_match = re.search(r'\(\d{4}\)\s+[\u2018\']([^\u2019\']{15,})[\u2019\']', text)
        if harvard_match:
            return harvard_match.group(1).strip()

    # --- Strategy 4: Vancouver/AMA — segments before year ---
    if year:
        year_str = str(year)
        year_idx = text.find(year_str)
        if year_idx > 10:
            before_year = text[:year_idx]
            chunks = [c.strip() for c in before_year.split(". ") if c.strip()]
            if len(chunks) >= 3:
                return chunks[-2].strip()
            elif len(chunks) == 2:
                return chunks[-1].strip()

    # --- Strategy 5: Fallback — after first ". " ---
    parts = text.split(". ", 2)
    if len(parts) >= 2:
        candidate = parts[1].strip().rstrip(".")
        if len(candidate) > 10:
            if year:
                year_str = str(year)
                if year_str in candidate:
                    candidate = candidate[:candidate.find(year_str)].strip().rstrip(".")
            return candidate

    return None


def extract_first_author(text):
    """Extract first author family name — format-agnostic.

    Handles:
      - Vancouver/AMA: "Caplin ME" → family="Caplin"
      - APA: "Caplin, M. E." → family="Caplin"
      - MLA/Chicago: "Caplin, M. E., et al." → family="Caplin"
      - Harvard: "Caplin, M.E." → family="Caplin"
    """
    # Get text before "et al" or first period
    author_block = text
    et_al_idx = text.lower().find("et al")
    if et_al_idx > 0:
        author_block = text[:et_al_idx].strip().rstrip(",.").strip()
    else:
        period_idx = text.find(". ")
        if period_idx > 0:
            author_block = text[:period_idx].strip()

    # First author is before the first comma
    first_author = author_block.split(",")[0].strip()

    # Extract family name: first word that's not all-caps initials
    parts = first_author.split()
    if not parts:
        return None

    family = parts[0]
    for part in parts:
        # Skip tokens like "ME", "AB" (1-3 uppercase letters = initials)
        if not re.match(r'^[A-Z]{1,3}$', part):
            family = part
            break

    return {"family": family, "full": first_author}


def extract_journal(text, year=None):
    """Extract journal name — best effort, format-agnostic."""
    if not year:
        return None

    year_str = str(year)
    year_idx = text.find(year_str)
    if year_idx < 0:
        return None

    # Journal is usually the segment just before the year
    before_year = text[:year_idx].strip().rstrip(".")
    chunks = before_year.rsplit(". ", 1)
    if len(chunks) == 2 and len(chunks[-1]) > 2:
        return chunks[-1].strip()
    return None


# ============================================================
# Main parser
# ============================================================

def parse_reference(num, text):
    """Parse a single reference into structured fields — format-agnostic."""
    result = {"num": num, "raw": text}

    # 1. DOI (format-independent)
    doi = extract_doi(text)
    if doi:
        result["doi"] = doi

    # 2. Year (format-independent)
    year = extract_year(text)
    if year:
        result["year"] = year

    # 3. Title (multi-strategy)
    title = extract_title(text, year)
    if title:
        result["title"] = title

    # 4. First author
    author = extract_first_author(text)
    if author:
        result["first_author_family"] = author["family"]
        result["first_author_full"] = author["full"]

    # 5. Journal (best effort)
    journal = extract_journal(text, year)
    if journal:
        result["journal"] = journal

    return result


# ============================================================
# Input readers
# ============================================================

REFERENCE_HEADINGS = {"references", "reference", "bibliography", "citations",
                      "works cited", "文献", "参考文献"}


def read_from_docx(docx_path):
    """Read .docx, locate References section, return list of {num, text}."""
    doc = Document(str(docx_path))
    ref_start_idx = None

    # Try Heading style first
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.lower() in REFERENCE_HEADINGS and para.style.name.startswith("Heading"):
            ref_start_idx = i
            break

    # Fallback: plain text match
    if ref_start_idx is None:
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip().lower() in REFERENCE_HEADINGS:
                ref_start_idx = i
                break

    if ref_start_idx is None:
        print("ERROR: Could not locate References section in .docx", file=sys.stderr)
        sys.exit(1)

    print(f"References section found at paragraph index {ref_start_idx}", file=sys.stderr)

    raw_refs = []
    i = ref_start_idx + 1
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        if not text:
            i += 1
            continue
        match = re.match(r'^(\d+)\.\s*(.+)', text)
        if match:
            ref_num = int(match.group(1))
            ref_text = match.group(2).strip()
            raw_refs.append({"num": ref_num, "text": ref_text})
        else:
            # Unnumbered paragraph — assign sequential number
            raw_refs.append({"num": len(raw_refs) + 1, "text": text})
        i += 1

    return raw_refs


def read_from_txt(txt_path):
    """Read .txt, each line = one reference (numbered or unnumbered)."""
    raw_refs = []
    with open(txt_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            match = re.match(r'^(\d+)\.\s*(.+)', line)
            if match:
                raw_refs.append({"num": int(match.group(1)), "text": match.group(2).strip()})
            else:
                raw_refs.append({"num": len(raw_refs) + 1, "text": line})
    return raw_refs


def read_from_json(json_path):
    """Read pre-extracted JSON (skip parsing)."""
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


# ============================================================
# Main
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Extract references from .docx, .txt, or .json into structured JSON. "
                    "Format-agnostic: supports Vancouver, APA, AMA, MLA, Chicago, Harvard, and unstructured text."
    )
    parser.add_argument("--input", "-i", required=True,
                        help="Input file: .docx (manuscript), .txt (reference list), or .json (pre-extracted)")
    parser.add_argument("--output", "-o", default=None,
                        help="Output JSON path (default: references_extracted.json in input dir)")
    parser.add_argument("--output-dir", "-d", default=None,
                        help="Output directory (overrides --output filename, keeps default name)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    # Resolve output path
    if args.output:
        output_path = Path(args.output)
    elif args.output_dir:
        output_path = Path(args.output_dir) / "references_extracted.json"
    else:
        output_path = input_path.parent / "references_extracted.json"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Read raw references based on file type
    ext = input_path.suffix.lower()
    if ext == ".docx":
        raw_refs = read_from_docx(input_path)
    elif ext == ".txt":
        raw_refs = read_from_txt(input_path)
    elif ext == ".json":
        raw_refs = read_from_json(input_path)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(raw_refs, f, ensure_ascii=False, indent=2)
        print(f"Pre-extracted JSON copied to: {output_path}", file=sys.stderr)
        print(f"Total references: {len(raw_refs)}", file=sys.stderr)
        return
    else:
        print(f"ERROR: Unsupported file type: {ext} (use .docx, .txt, or .json)", file=sys.stderr)
        sys.exit(1)

    if not raw_refs:
        print("WARNING: No references found in input file.", file=sys.stderr)

    # Parse each reference
    parsed_refs = []
    for ref in raw_refs:
        parsed = parse_reference(ref["num"], ref["text"])
        parsed_refs.append(parsed)
        print(f"  [{parsed['num']}] Author: {parsed.get('first_author_family', '?')} | "
              f"Year: {parsed.get('year', '?')} | "
              f"DOI: {parsed.get('doi', '—')} | "
              f"Title: {parsed.get('title', '?')[:60]}...", file=sys.stderr)

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(parsed_refs, f, ensure_ascii=False, indent=2)

    has_doi = sum(1 for r in parsed_refs if r.get("doi"))
    print(f"\nOutput: {output_path}", file=sys.stderr)
    print(f"Total references extracted: {len(parsed_refs)}", file=sys.stderr)
    print(f"  With DOI: {has_doi}", file=sys.stderr)
    print(f"  Without DOI: {len(parsed_refs) - has_doi}", file=sys.stderr)


if __name__ == "__main__":
    main()
